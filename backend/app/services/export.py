"""Export appeal letters as PDF or DOC."""

import io
from datetime import date
from html import unescape
from xml.sax.saxutils import escape

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
)

from ..models import Appeal


def _strip_html(text: str) -> str:
    import re

    text = re.sub(r"<[^>]+>", "", text)
    text = unescape(text)
    return text.strip()


def _build_letter_header(appeal: Appeal) -> dict:
    """Extract metadata for the letter header from the appeal chain."""
    claim = appeal.denial.claim
    practice = claim.practice
    payer = claim.payer
    patient = claim.patient
    denial = appeal.denial

    today = date.today()

    # City, ST ZIP from whatever profile fields are filled.
    city_state = ", ".join(p for p in (practice.city, practice.state) if p)
    if practice.zip_code:
        city_state = f"{city_state} {practice.zip_code}".strip()

    # Signature: "Sarah Chen, MD" if available, else fall back to the practice.
    provider = practice.primary_provider_name or practice.name
    if practice.primary_provider_name and practice.primary_provider_credentials:
        provider = (
            f"{practice.primary_provider_name}, "
            f"{practice.primary_provider_credentials}"
        )

    ids = []
    if practice.npi:
        ids.append(f"NPI: {practice.npi}")
    if practice.tax_id:
        ids.append(f"Tax ID: {practice.tax_id}")

    return {
        "practice_name": practice.name,
        "practice_address": practice.address_line1 or "",
        "practice_address2": practice.address_line2 or "",
        "practice_city_state_zip": city_state,
        "practice_phone": practice.phone or "",
        "practice_fax": practice.fax or "",
        "practice_ids": " | ".join(ids),
        "provider_signature": provider,
        "date": today.strftime("%B %d, %Y"),
        "payer_name": payer.name,
        "payer_address": "Claims Department",
        "payer_city_state_zip": "",
        "patient_name": patient.name,
        "patient_dob": patient.dob.strftime("%m/%d/%Y") if patient.dob else "",
        "claim_dos": claim.date_of_service.strftime("%B %d, %Y") if claim.date_of_service else "",
        "claim_cpt": ", ".join(claim.cpt_codes) if claim.cpt_codes else "",
        "claim_icd": ", ".join(claim.icd_codes) if claim.icd_codes else "",
        "denial_code": denial.denial_code or "",
        "billed_amount": f"${denial.denied_amount:,.2f}" if denial.denied_amount else "",
    }


def generate_pdf(appeal: Appeal) -> bytes:
    """Generate a professional appeal letter as PDF using reportlab."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        topMargin=1.2 * inch,
        bottomMargin=1 * inch,
    )

    styles = getSampleStyleSheet()
    letter_body = ParagraphStyle(
        "LetterBody",
        parent=styles["Normal"],
        fontSize=11,
        leading=15,
        spaceAfter=8,
        fontName="Times-Roman",
    )
    letter_header = ParagraphStyle(
        "LetterHeader",
        parent=styles["Normal"],
        fontSize=12,
        leading=16,
        fontName="Times-Bold",
    )

    meta = _build_letter_header(appeal)

    elements: list = []

    # ReportLab parses Paragraph input as XML-like markup, so any dynamic value
    # containing &, <, or > (e.g. "Blue Cross & Blue Shield") must be escaped to
    # avoid a parse error; only the literal <b>...</b> tags below stay unescaped.

    # Sender (omit any letterhead lines the practice hasn't filled in)
    elements.append(
        Paragraph(f"<b>{escape(meta['practice_name'])}</b>", letter_header)
    )
    if meta["practice_address"]:
        elements.append(Paragraph(escape(meta["practice_address"]), letter_body))
    if meta["practice_address2"]:
        elements.append(Paragraph(escape(meta["practice_address2"]), letter_body))
    if meta["practice_city_state_zip"]:
        elements.append(
            Paragraph(escape(meta["practice_city_state_zip"]), letter_body)
        )
    contact = " · ".join(
        p
        for p in (
            f"Phone: {meta['practice_phone']}" if meta["practice_phone"] else "",
            f"Fax: {meta['practice_fax']}" if meta["practice_fax"] else "",
        )
        if p
    )
    if contact:
        elements.append(Paragraph(escape(contact), letter_body))
    if meta["practice_ids"]:
        elements.append(Paragraph(escape(meta["practice_ids"]), letter_body))
    elements.append(Spacer(1, 18))

    # Date
    elements.append(Paragraph(escape(meta["date"]), letter_body))
    elements.append(Spacer(1, 12))

    # Recipient
    elements.append(
        Paragraph(f"<b>{escape(meta['payer_name'])}</b>", letter_header)
    )
    elements.append(Paragraph(escape(meta["payer_address"]), letter_body))
    if meta["payer_city_state_zip"]:
        elements.append(Paragraph(escape(meta["payer_city_state_zip"]), letter_body))
    elements.append(Spacer(1, 18))

    # Subject
    elements.append(
        Paragraph(
            f"<b>RE: Appeal of Denied Claim — Patient: {escape(meta['patient_name'])}"
            f" | Date of Service: {escape(meta['claim_dos'])}"
            f" | Denial Code: {escape(meta['denial_code'])}</b>",
            letter_body,
        )
    )
    elements.append(Spacer(1, 14))

    # Salutation
    elements.append(Paragraph("To Whom It May Concern:", letter_body))
    elements.append(Spacer(1, 10))

    # Letter body — strip HTML tags, use plain paragraphs
    letter_text = appeal.letter_text or ""
    clean = _strip_html(letter_text)
    for para in clean.split("\n"):
        stripped = para.strip()
        if stripped:
            elements.append(Paragraph(escape(stripped), letter_body))
    elements.append(Spacer(1, 14))

    # Closing
    elements.append(Paragraph("Sincerely,", letter_body))
    elements.append(Spacer(1, 28))
    elements.append(Paragraph(escape(meta["provider_signature"]), letter_body))
    if meta["provider_signature"] != meta["practice_name"]:
        elements.append(Paragraph(escape(meta["practice_name"]), letter_body))

    doc.build(elements)
    buf.seek(0)
    return buf.read()


def generate_doc(appeal: Appeal) -> bytes:
    """Generate a professional appeal letter as DOCX using python-docx."""
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    meta = _build_letter_header(appeal)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    section = doc.sections[0]
    section.top_margin = Inches(1.2)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    def add_para(text, bold=False, size=11, space_before=0, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        run.bold = bold
        return p

    # Header (omit letterhead lines the practice hasn't filled in)
    add_para(meta["practice_name"], bold=True, size=12, space_after=2)
    if meta["practice_address"]:
        add_para(meta["practice_address"], size=11, space_after=2)
    if meta["practice_address2"]:
        add_para(meta["practice_address2"], size=11, space_after=2)
    if meta["practice_city_state_zip"]:
        add_para(meta["practice_city_state_zip"], size=11, space_after=2)
    contact = " · ".join(
        p
        for p in (
            f"Phone: {meta['practice_phone']}" if meta["practice_phone"] else "",
            f"Fax: {meta['practice_fax']}" if meta["practice_fax"] else "",
        )
        if p
    )
    if contact:
        add_para(contact, size=11, space_after=2)
    if meta["practice_ids"]:
        add_para(meta["practice_ids"], size=11, space_after=2)
    add_para("", size=4, space_after=12)  # spacer before date

    add_para(meta["date"], size=11, space_after=12)

    # Recipient
    add_para(meta["payer_name"], bold=True, size=12, space_after=2)
    add_para(meta["payer_address"], size=11, space_after=2)
    if meta["payer_city_state_zip"]:
        add_para(meta["payer_city_state_zip"], size=11, space_after=18)

    # Subject
    subject = (
        f"RE: Appeal of Denied Claim — Patient: {meta['patient_name']}"
        f" | Date of Service: {meta['claim_dos']}"
        f" | Denial Code: {meta['denial_code']}"
    )
    add_para(subject, bold=True, size=11, space_after=14)

    add_para("To Whom It May Concern:", size=11, space_after=10)

    # Letter body
    letter_text = appeal.letter_text or ""
    clean = _strip_html(letter_text)
    for para in clean.split("\n"):
        stripped = para.strip()
        if stripped:
            add_para(stripped, size=11, space_after=8)

    add_para("", size=11, space_after=14)
    add_para("Sincerely,", size=11, space_after=28)
    add_para(meta["provider_signature"], size=11, space_after=2)
    if meta["provider_signature"] != meta["practice_name"]:
        add_para(meta["practice_name"], size=11, space_after=2)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
